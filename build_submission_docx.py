"""Build Molecular Autism submission .docx files with proper Word table objects."""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re


def set_double_spacing(doc):
    """Set double spacing for all paragraphs."""
    for para in doc.paragraphs:
        para.paragraph_format.line_spacing = 2.0


def add_line_numbers(doc):
    """Add continuous line numbering to document."""
    sectPr = doc.sections[0]._sectPr
    lnNumType = sectPr.makeelement(qn('w:lnNumType'), {
        qn('w:countBy'): '1',
        qn('w:restart'): 'continuous',
    })
    sectPr.append(lnNumType)


def add_page_numbers(doc):
    """Add page numbers to footer."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    fldChar1 = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run._element.append(fldChar1)
    run2 = para.add_run()
    instrText = run2._element.makeelement(qn('w:instrText'), {})
    instrText.text = ' PAGE '
    run2._element.append(instrText)
    run3 = para.add_run()
    fldChar2 = run3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run3._element.append(fldChar2)


def make_table(doc, headers, rows, title=None, legend=None):
    """Create a proper Word table object."""
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    if legend:
        p = doc.add_paragraph()
        run = p.add_run(legend)
        run.font.size = Pt(9)
        run.italic = True

    doc.add_paragraph()  # spacing


def build_main_manuscript():
    """Build the main manuscript with proper Word tables."""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Genome-Wide Stability Profiling Identifies a Mitochondrial Oxidative Phosphorylation Cluster in Autism Spectrum Disorder Brain Transcriptomics')
    run.bold = True
    run.font.size = Pt(14)

    # Author
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Ray Sigmon').bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Alpha Research Labs, Northwest Arkansas, USA')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Correspondence: alphalabsincorp@gmail.com')

    # Keywords
    p = doc.add_paragraph()
    run = p.add_run('Keywords: ')
    run.bold = True
    p.add_run('autism spectrum disorder, mitochondrial dysfunction, oxidative phosphorylation, transcriptomics, stability profiling, cell danger response, gene co-expression, purinergic signaling, postmortem brain, meta-analysis')

    # Now read the markdown and insert text sections, replacing tables with Word tables
    with open('/home/kai001/riker-engine/ASD_Molecular_Autism_Submission.md', 'r') as f:
        content = f.read()

    # Split into sections by ## headers
    # We'll process the text manually, skipping the title/keywords (already added)
    # and inserting proper tables where needed

    lines = content.split('\n')

    # Find where Abstract starts
    in_table = False
    table_lines = []
    skip_until_next_section = False
    i = 0

    # Skip to Abstract
    while i < len(lines) and not lines[i].startswith('## Abstract'):
        i += 1

    while i < len(lines):
        line = lines[i]

        # Handle section headers
        if line.startswith('## '):
            heading = line.replace('## ', '').strip()
            if heading in ['Additional files']:
                # Process Additional files section
                doc.add_heading(heading, level=2)
                i += 1
                while i < len(lines):
                    if lines[i].strip():
                        doc.add_paragraph(lines[i].strip())
                    i += 1
                break
            doc.add_heading(heading, level=2)
            i += 1
            continue

        if line.startswith('### '):
            heading = line.replace('### ', '').strip()
            doc.add_heading(heading, level=3)
            i += 1
            continue

        # Skip the markdown horizontal rules
        if line.strip() == '---':
            i += 1
            continue

        # Detect table starts (markdown pipe tables)
        if '|' in line and line.strip().startswith('|'):
            # Collect all table lines
            table_lines = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1

            # Parse the table
            if len(table_lines) >= 3:  # header + separator + at least one row
                headers = [c.strip() for c in table_lines[0].split('|')[1:-1]]
                rows = []
                for tl in table_lines[2:]:  # skip separator line
                    row = [c.strip() for c in tl.split('|')[1:-1]]
                    if len(row) == len(headers):
                        rows.append(row)

                make_table(doc, headers, rows)
            continue

        # Handle bullet points in Declarations
        if line.strip().startswith('- **'):
            text = line.strip()[2:]  # Remove "- "
            p = doc.add_paragraph(style='List Bullet')
            # Parse bold parts
            parts = text.split('**')
            for idx, part in enumerate(parts):
                if idx % 2 == 1:  # odd indices are bold
                    p.add_run(part).bold = True
                else:
                    p.add_run(part)
            i += 1
            continue

        # Regular paragraph text
        if line.strip():
            text = line.strip()
            # Handle bold markers in text
            p = doc.add_paragraph()
            # Simple bold handling for **text**
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        i += 1

    # Apply formatting
    set_double_spacing(doc)
    add_line_numbers(doc)
    add_page_numbers(doc)

    doc.save('/home/kai001/Downloads/ASD_Molecular_Autism_Submission_final.docx')
    print("Main manuscript saved")


def build_additional_file_2():
    """Build Additional file 2 with proper Word tables."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    doc.add_heading('Additional file 2', level=1)

    doc.add_paragraph(
        'Inverse-variance weighted random-effects meta-analysis across three brain discovery cohorts '
        'for 41 mitochondrial and energy metabolism cluster genes.'
    ).runs[0].bold = True

    doc.add_paragraph(
        'Meta-analysis was computed using DerSimonian-Laird random-effects estimation (REML fallback '
        'when k \u2265 3) for genes surviving Phase 5 replication. Genes eliminated at Phase 5 due to '
        'significant opposite-direction expression in blood replication cohorts are noted with the '
        'eliminating cohort. This brain-blood directional divergence is consistent with expected '
        'tissue-specific expression for brain mitochondrial genes. RE = random-effects pooled estimate '
        '(log2FC); SE = standard error; 95% CI = RE \u00b1 1.96 \u00d7 SE; I\u00b2 = between-study '
        'heterogeneity. Mean log2FC = unweighted cross-dataset average from Phase 1, reported for all genes.'
    )

    # Read the corrected markdown for table data
    with open('/home/kai001/riker-engine/Additional_file_2_meta_analysis.md', 'r') as f:
        content = f.read()

    headers = ['Gene', 'Mean log2FC', 'Direction', 'Phase 5 Status', 'RE', 'SE', '95% CI', 'p', 'I\u00b2 (%)']

    # Parse tables from the markdown
    lines = content.split('\n')

    current_section = None
    table_rows = []

    for line in lines:
        if line.startswith('### 26-gene'):
            if table_rows:
                make_table(doc, headers, table_rows, current_section)
                table_rows = []
            current_section = '26-gene mitochondrial/OxPhos core'
        elif line.startswith('### Additional 15'):
            if table_rows:
                make_table(doc, headers, table_rows, current_section)
                table_rows = []
            current_section = 'Additional 15 energy metabolism genes'
        elif line.startswith('### Summary'):
            if table_rows:
                make_table(doc, headers, table_rows, current_section)
                table_rows = []
            current_section = None
        elif '|' in line and line.strip().startswith('|') and not line.strip().startswith('| Gene') and not '---' in line:
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if len(cells) == len(headers):
                # Replace em-dash with regular dash for consistency
                cells = [c.replace('\u2014', '\u2014') for c in cells]
                table_rows.append(cells)

    if table_rows:
        make_table(doc, headers, table_rows, current_section)

    # Add summary section
    doc.add_heading('Summary', level=3)

    # Parse summary from markdown
    in_summary = False
    for line in lines:
        if line.startswith('### Summary'):
            in_summary = True
            continue
        if in_summary and line.startswith('### '):
            break
        if in_summary and line.strip().startswith('- '):
            doc.add_paragraph(line.strip()[2:], style='List Bullet')
        elif in_summary and line.strip().startswith('Note'):
            doc.add_paragraph(line.strip())
        elif in_summary and line.strip() and not line.startswith('Of ') and not line.startswith('-'):
            if line.strip().startswith('Of '):
                doc.add_paragraph(line.strip())

    # Manually add the corrected summary text to be sure
    # Clear and re-add since parsing might miss the structure
    doc.add_paragraph(
        'Of 41 mitochondrial and energy metabolism cluster genes:'
    )

    summary_bullets = [
        '19 survived Phase 5 replication and entered Phase 6 meta-analysis',
        '20 were eliminated at Phase 5 due to significant opposite-direction expression in blood replication cohorts (12 eliminated by GSE18123 blood LCL, 7 by GSE26415 blood leukocytes, 1 by GSE42133 blood leukocytes)',
        '2 were not in the core gene set for the representative run (COX7A1 at 47/50 stability, SLC25A12 at 48/50 stability \u2014 both iron-clad overall but not in run 001)',
    ]
    for b in summary_bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_paragraph('Of the 19 genes with meta-analysis data:')

    meta_bullets = [
        '8 showed significant pooled random-effects (p < 0.05): CYBA, GPI, IDH3A, OGDHL, SLC25A27, NDUFAF5, PFKP, HK2',
        '3 showed significant effects with moderate-to-high heterogeneity: ME3, UQCRC1, HK2',
        '6 showed non-significant effects with high heterogeneity (I\u00b2 > 80%): ATP5F1A (90.7%), PFKM (87.6%), NFS1 (86.7%), PYGB (84.9%), ALDOC (82.9%), FH (82.6%)',
    ]
    for b in meta_bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_paragraph(
        'Note on sign discrepancies: The Phase 6 random-effects estimate uses inverse-variance weighting '
        'across the three brain discovery cohorts and may differ in magnitude or sign from the Phase 1 '
        'unweighted mean log2FC. For example, ME3 has a Phase 1 mean log2FC of +0.08 (up) but a Phase 6 '
        'RE of \u22120.296 (down, p = 0.025). This occurs because the three cohorts contribute different '
        'effect sizes and the inverse-variance weighting can shift the pooled estimate relative to a simple '
        'average. Similar sign discrepancies appear for ATP5F1A, PFKM, SLC25A3, ALDOC, and FH. These '
        'discrepancies reflect between-cohort heterogeneity in effect magnitude across different platforms '
        'and brain regions \u2014 not errors in either estimate. The stability finding (consistent '
        'co-clustering in \u226590% of 50 runs) is methodologically distinct from effect size homogeneity.'
    )

    set_double_spacing(doc)
    add_line_numbers(doc)
    add_page_numbers(doc)

    doc.save('/home/kai001/Downloads/Additional_file_2_meta_analysis_final.docx')
    print("Additional file 2 saved")


def build_cover_letter():
    """Build cover letter with proper formatting."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    doc.add_paragraph('Dear Editors,')
    doc.add_paragraph('')

    doc.add_paragraph(
        'I am submitting the enclosed manuscript, "Genome-Wide Stability Profiling Identifies a '
        'Mitochondrial Oxidative Phosphorylation Cluster in Autism Spectrum Disorder Brain '
        'Transcriptomics," for consideration as a Research article in Molecular Autism.'
    )

    doc.add_paragraph(
        'This study applies a novel computational approach \u2014 genome-wide stability profiling across '
        '50 independent pipeline runs \u2014 to ASD postmortem brain transcriptomics. The central finding '
        'is a 26-gene mitochondrial and energy metabolism cluster (expanding to 41 genes) that achieves '
        'iron-clad stability (\u226590% cross-run reproducibility) in a blind, hypothesis-free analysis '
        'across three discovery cohorts and 664 total samples. This mitochondrial signature is consistent '
        'with the Cell Danger Response hypothesis and complements existing metabolomic evidence of '
        'mitochondrial dysfunction in ASD. Of the 376 iron-clad genes identified, 352 are not in the '
        'SFARI database, suggesting substantial convergent biology beyond established candidate genes.'
    )

    doc.add_paragraph('I believe this manuscript is appropriate for Molecular Autism because it:')

    bullets = [
        'Identifies convergent molecular pathobiology (mitochondrial dysfunction) across independent ASD cohorts using a reproducible, openly available methodology',
        'Bridges transcriptomic and metabolomic levels of analysis, relevant to understanding ASD etiology',
        'Reports all findings with stability quantification, effect size meta-analysis with confidence intervals, and explicit platform limitations, consistent with the journal\u2019s emphasis on balanced reporting and rigorous limitations disclosure',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_paragraph(
        'All source code, data, and stability reports are publicly available on GitHub and archived in '
        'Zenodo (https://doi.org/10.5281/zenodo.19623672). The analysis is fully reproducible from a '
        'single master seed.'
    )

    doc.add_paragraph(
        'I am the sole author and have approved the manuscript for submission. The content has not been '
        'published or submitted for publication elsewhere. I declare no competing interests.'
    )

    doc.add_paragraph(
        'I respectfully request consideration for an APC waiver. This work was conducted independently '
        'without institutional affiliation or external funding. I am an independent computational '
        'researcher with no grant support or institutional APC coverage. I am committed to open science '
        'and have made all code and data freely available under an AGPL-3.0 license.'
    )

    doc.add_paragraph('Thank you for considering this submission.')
    doc.add_paragraph('')
    doc.add_paragraph('Sincerely,')
    doc.add_paragraph('Ray Sigmon')
    doc.add_paragraph('Alpha Research Labs')
    doc.add_paragraph('Northwest Arkansas, USA')
    doc.add_paragraph('alphalabsincorp@gmail.com')

    set_double_spacing(doc)

    doc.save('/home/kai001/Downloads/ASD_Molecular_Autism_Cover_Letter_final.docx')
    print("Cover letter saved")


if __name__ == '__main__':
    build_main_manuscript()
    build_additional_file_2()
    build_cover_letter()
